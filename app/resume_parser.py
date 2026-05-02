from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader


class UnsupportedFileTypeError(ValueError):
    """Raised when the uploaded resume format is not supported."""


async def extract_text_from_upload(file: UploadFile) -> str:
    filename = file.filename or ""
    extension = Path(filename).suffix.lower()
    content = await file.read()

    if not content:
        raise ValueError("Uploaded file is empty.")

    if extension == ".pdf":
        text = _extract_pdf_text(content)
    elif extension == ".docx":
        text = _extract_docx_text(content)
    elif extension == ".txt":
        text = _extract_txt_text(content)
    else:
        raise UnsupportedFileTypeError(
            "Unsupported file type. Please upload a PDF, DOCX, or TXT resume."
        )

    if not text.strip():
        raise ValueError("Could not read text from the uploaded resume.")

    return text


def _extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                table_cells.append(cell.text)

    return "\n".join(paragraphs + table_cells)


def _extract_txt_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")

